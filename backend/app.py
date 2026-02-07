from flask import Flask, request, jsonify, send_file, Response, stream_with_context
from flask_cors import CORS
from pdf2docx import Converter
import os
import uuid
from werkzeug.utils import secure_filename
from ai_service import ai_service
import fitz
from docx import Document

app = Flask(__name__)
CORS(app)

UPLOAD_FOLDER = 'uploads'
DOWNLOAD_FOLDER = 'downloads'
ALLOWED_EXTENSIONS = {'pdf'}

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['DOWNLOAD_FOLDER'] = DOWNLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/api/convert', methods=['POST'])
def convert_pdf():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Only PDF files are allowed'}), 400
    
    try:
        filename = secure_filename(file.filename)
        original_name = os.path.splitext(filename)[0]
        unique_id = str(uuid.uuid4())
        pdf_filename = f'{unique_id}_{filename}'
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
        
        docx_filename = f'{unique_id}_{original_name}.docx'
        docx_path = os.path.join(app.config['DOWNLOAD_FOLDER'], docx_filename)
        
        file.save(pdf_path)
        
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        
        os.remove(pdf_path)
        
        download_filename = f'{original_name}.docx'
        return jsonify({
            'success': True,
            'downloadUrl': f'/api/download/{docx_filename}',
            'filename': download_filename
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/extract-text', methods=['POST'])
def extract_pdf_text():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Only PDF files are allowed'}), 400
    
    try:
        filename = secure_filename(file.filename)
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        file.save(pdf_path)
        
        doc = fitz.open(pdf_path)
        text_content = ''
        for page in doc:
            text_content += page.get_text()
        doc.close()
        
        os.remove(pdf_path)
        
        return jsonify({
            'success': True,
            'text': text_content,
            'char_count': len(text_content)
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    try:
        file_path = os.path.join(app.config['DOWNLOAD_FOLDER'], filename)
        download_name = request.args.get('name', filename)
        return send_file(file_path, as_attachment=True, download_name=download_name)
    except Exception as e:
        return jsonify({'error': str(e)}), 404

@app.route('/api/health', methods=['GET'])
def health_check():
    return jsonify({'status': 'ok'})

@app.route('/api/ai/status', methods=['GET'])
def ai_status():
    return jsonify({
        'enabled': ai_service.is_enabled(),
        'features': {
            'summary': ai_service.enable_summary,
            'translate': ai_service.enable_translate,
            'chat': ai_service.enable_chat
        }
    })

@app.route('/api/ai/summary', methods=['POST'])
def generate_summary():
    data = request.get_json()
    text = data.get('text', '')
    api_key = data.get('api_key')
    
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    
    try:
        import asyncio
        summary = asyncio.run(ai_service.generate_summary(text, api_key))
        
        if summary:
            return jsonify({
                'success': True,
                'summary': summary
            })
        else:
            return jsonify({'error': 'Failed to generate summary'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ai/translate', methods=['POST'])
def translate_text():
    data = request.get_json()
    text = data.get('text', '')
    target_lang = data.get('target_lang', '中文')
    api_key = data.get('api_key')
    
    if not text:
        return jsonify({'error': 'No text provided'}), 400
    
    try:
        import asyncio
        translation = asyncio.run(ai_service.translate_text(text, target_lang, api_key))
        
        if translation:
            return jsonify({
                'success': True,
                'translation': translation
            })
        else:
            return jsonify({'error': 'Failed to translate'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ai/chat', methods=['POST'])
def chat_with_document():
    data = request.get_json()
    question = data.get('question', '')
    context = data.get('context', '')
    api_key = data.get('api_key')
    
    if not question:
        return jsonify({'error': 'No question provided'}), 400
    
    try:
        import asyncio
        answer = asyncio.run(ai_service.chat_with_document(question, context, api_key))
        
        if answer:
            return jsonify({
                'success': True,
                'answer': answer
            })
        else:
            return jsonify({'error': 'Failed to get answer'}), 500
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/ai/chat/stream', methods=['POST'])
def chat_with_document_stream():
    data = request.get_json()
    question = data.get('question', '')
    context = data.get('context', '')
    api_key = data.get('api_key')
    
    if not question:
        return jsonify({'error': 'No question provided'}), 400
    
    def generate():
        try:
            stream_response = ai_service.chat_with_document_stream(question, context, api_key)
            if stream_response is None:
                yield 'data: {"error": "Failed to get stream response"}\n\n'
                return
            
            for chunk in stream_response:
                if chunk.choices and len(chunk.choices) > 0:
                    delta = chunk.choices[0].delta
                    if hasattr(delta, 'content') and delta.content:
                        yield f'data: {{"content": "{delta.content}"}}\n\n'
            
            yield 'data: {"done": true}\n\n'
        except Exception as e:
            yield f'data: {{"error": "{str(e)}"}}\n\n'
    
    return Response(stream_with_context(generate()), mimetype='text/event-stream')

@app.route('/api/ai/test-key', methods=['POST'])
def test_api_key():
    data = request.get_json()
    api_key = data.get('api_key')
    
    if not api_key:
        return jsonify({'error': 'No API key provided'}), 400
    
    result = ai_service.test_api_key(api_key)
    
    if result['valid']:
        return jsonify({
            'success': True,
            'message': result['message']
        })
    else:
        return jsonify({
            'success': False,
            'error': result['error']
        }), 400

@app.route('/api/translate-pdf', methods=['POST'])
def translate_pdf_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    target_lang = request.form.get('target_lang', '中文')
    api_key = request.form.get('api_key')
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Only PDF files are allowed'}), 400
    
    try:
        filename = secure_filename(file.filename)
        original_name = os.path.splitext(filename)[0]
        unique_id = str(uuid.uuid4())
        pdf_filename = f'{unique_id}_{filename}'
        pdf_path = os.path.join(app.config['UPLOAD_FOLDER'], pdf_filename)
        
        file.save(pdf_path)
        
        docx_filename = f'{unique_id}_{original_name}.docx'
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
        
        cv = Converter(pdf_path)
        cv.convert(docx_path)
        cv.close()
        
        os.remove(pdf_path)
        
        doc = Document(docx_path)
        
        translated_count = 0
        failed_count = 0
        
        for para in doc.paragraphs:
            if para.text.strip():
                try:
                    import asyncio
                    original_text = para.text
                    translated_text = asyncio.run(ai_service.translate_text(original_text, target_lang, api_key))
                    if translated_text and translated_text.strip():
                        for run in para.runs:
                            if run.text.strip():
                                run.text = translated_text
                                translated_count += 1
                                break
                except Exception as e:
                    print(f'翻译段落失败: {str(e)}')
                    failed_count += 1
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            try:
                                import asyncio
                                original_text = para.text
                                translated_text = asyncio.run(ai_service.translate_text(original_text, target_lang, api_key))
                                if translated_text and translated_text.strip():
                                    for run in para.runs:
                                        if run.text.strip():
                                            run.text = translated_text
                                            translated_count += 1
                                            break
                            except Exception as e:
                                print(f'翻译表格单元格失败: {str(e)}')
                                failed_count += 1
        
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        if para.text.strip():
                            try:
                                import asyncio
                                original_text = para.text
                                translated_text = asyncio.run(ai_service.translate_text(original_text, target_lang, api_key))
                                if translated_text and translated_text.strip():
                                    for run in para.runs:
                                        if run.text.strip():
                                            run.text = translated_text
                                            translated_count += 1
                                            break
                            except Exception as e:
                                print(f'翻译页眉失败: {str(e)}')
                                failed_count += 1
            
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        if para.text.strip():
                            try:
                                import asyncio
                                original_text = para.text
                                translated_text = asyncio.run(ai_service.translate_text(original_text, target_lang, api_key))
                                if translated_text and translated_text.strip():
                                    for run in para.runs:
                                        if run.text.strip():
                                            run.text = translated_text
                                            translated_count += 1
                                            break
                            except Exception as e:
                                print(f'翻译页脚失败: {str(e)}')
                                failed_count += 1
        
        print(f'翻译完成: 成功 {translated_count} 处, 失败 {failed_count} 处')
        
        translated_filename = f'{unique_id}_{original_name}_translated.docx'
        translated_path = os.path.join(app.config['DOWNLOAD_FOLDER'], translated_filename)
        doc.save(translated_path)
        
        os.remove(docx_path)
        
        download_filename = f'{original_name}_translated.docx'
        return jsonify({
            'success': True,
            'downloadUrl': f'/api/download/{translated_filename}',
            'filename': download_filename
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/translate-word', methods=['POST'])
def translate_word_file():
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    target_lang = request.form.get('target_lang', '中文')
    api_key = request.form.get('api_key')
    
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not file.filename.lower().endswith('.docx'):
        return jsonify({'error': 'Only DOCX files are allowed'}), 400
    
    try:
        filename = secure_filename(file.filename)
        original_name = os.path.splitext(filename)[0]
        unique_id = str(uuid.uuid4())
        docx_filename = f'{unique_id}_{filename}'
        docx_path = os.path.join(app.config['UPLOAD_FOLDER'], docx_filename)
        
        file.save(docx_path)
        
        doc = Document(docx_path)
        
        translated_count = 0
        failed_count = 0
        
        for para in doc.paragraphs:
            if para.text.strip():
                try:
                    import asyncio
                    original_text = para.text
                    translated_text = asyncio.run(ai_service.translate_text(original_text, target_lang, api_key))
                    if translated_text and translated_text.strip():
                        for run in para.runs:
                            if run.text.strip():
                                run.text = translated_text
                                translated_count += 1
                                break
                except Exception as e:
                    print(f'翻译段落失败: {str(e)}')
                    failed_count += 1
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            try:
                                import asyncio
                                original_text = para.text
                                translated_text = asyncio.run(ai_service.translate_text(original_text, target_lang, api_key))
                                if translated_text and translated_text.strip():
                                    for run in para.runs:
                                        if run.text.strip():
                                            run.text = translated_text
                                            translated_count += 1
                                            break
                            except Exception as e:
                                print(f'翻译表格单元格失败: {str(e)}')
                                failed_count += 1
        
        for section in doc.sections:
            for header in [section.header, section.first_page_header, section.even_page_header]:
                if header:
                    for para in header.paragraphs:
                        if para.text.strip():
                            try:
                                import asyncio
                                original_text = para.text
                                translated_text = asyncio.run(ai_service.translate_text(original_text, target_lang, api_key))
                                if translated_text and translated_text.strip():
                                    for run in para.runs:
                                        if run.text.strip():
                                            run.text = translated_text
                                            translated_count += 1
                                            break
                            except Exception as e:
                                print(f'翻译页眉失败: {str(e)}')
                                failed_count += 1
            
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                if footer:
                    for para in footer.paragraphs:
                        if para.text.strip():
                            try:
                                import asyncio
                                original_text = para.text
                                translated_text = asyncio.run(ai_service.translate_text(original_text, target_lang, api_key))
                                if translated_text and translated_text.strip():
                                    for run in para.runs:
                                        if run.text.strip():
                                            run.text = translated_text
                                            translated_count += 1
                                            break
                            except Exception as e:
                                print(f'翻译页脚失败: {str(e)}')
                                failed_count += 1
        
        print(f'翻译完成: 成功 {translated_count} 处, 失败 {failed_count} 处')
        
        translated_filename = f'{unique_id}_{original_name}_translated.docx'
        translated_path = os.path.join(app.config['DOWNLOAD_FOLDER'], translated_filename)
        doc.save(translated_path)
        
        os.remove(docx_path)
        
        download_filename = f'{original_name}_translated.docx'
        return jsonify({
            'success': True,
            'downloadUrl': f'/api/download/{translated_filename}',
            'filename': download_filename
        })
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)
